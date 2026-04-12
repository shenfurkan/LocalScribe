import csv
import html as html_lib
import io
import re
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


# ── Time formatters ────────────────────────────────────────────────────────────

def _carry_srt_time(h: int, m: int, s: int, ms: int) -> tuple[int, int, int, int]:
    """Cascade millisecond/second/minute overflows produced by rounding."""
    if ms >= 1000:
        s  += ms // 1000
        ms  = ms % 1000
    if s >= 60:
        m  += s // 60
        s   = s % 60
    if m >= 60:
        h  += m // 60
        m   = m % 60
    return h, m, s, ms


def _srt_time(seconds: float) -> str:
    """Converts float seconds to SRT timecode: HH:MM:SS,mmm"""
    h   = int(seconds // 3600)
    m   = int((seconds % 3600) // 60)
    s   = int(seconds % 60)
    ms  = int(round((seconds % 1) * 1000))
    h, m, s, ms = _carry_srt_time(h, m, s, ms)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _vtt_time(seconds: float) -> str:
    """Same as _srt_time but uses '.' separator (WebVTT spec)."""
    return _srt_time(seconds).replace(",", ".")


def _readable_time(seconds: float) -> str:
    """Converts to [HH:MM:SS] suitable for TXT/DOCX inline timestamps."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"[{h:02d}:{m:02d}:{s:02d}]"


# ── Smart SRT re-segmentation ──────────────────────────────────────────────────

_SENTENCE_END = re.compile(r'[.!?]["\']?$')


def resegment_for_srt(
    segments: list,
    max_words: int = 8,
    max_duration: float = 10.0,
    max_chars: int = 80,
    sentence_aware: bool = True,
) -> list:
    """
    Re-segments a transcript using word-level timestamps into subtitle-
    friendly chunks. Falls back to segment-level splitting when word
    timestamps are absent.

    Args:
        segments:       list of segment dicts (each may have a ``words`` list).
        max_words:      maximum words per subtitle card.
        max_duration:   maximum duration (seconds) per subtitle card.
        max_chars:      maximum characters per subtitle card.
        sentence_aware: when True, a sentence boundary always starts a new card.

    Returns:
        A new list of segment dicts compatible with ``export_srt``.
    """
    # Flatten all words from every segment into a single ordered stream.
    all_words: list[dict] = []
    for seg in segments:
        words = seg.get("words") or []
        if words:
            all_words.extend(words)
        else:
            # No word timestamps — treat the whole segment as one word-blob.
            all_words.append({
                "word":  seg["text"],
                "start": seg["start"],
                "end":   seg["end"],
            })

    if not all_words:
        return segments

    result: list[dict] = []
    bucket: list[str] = []
    bucket_start: float = all_words[0]["start"]
    bucket_end:   float = all_words[0]["end"]

    def _flush() -> None:
        nonlocal bucket, bucket_start, bucket_end
        text = " ".join(bucket).strip()
        if text:
            result.append({
                "start": bucket_start,
                "end":   bucket_end,
                "text":  text,
                "words": [],
            })
        bucket = []

    for word in all_words:
        word_text: str   = word.get("word", "").strip()
        if not word_text:
            continue

        w_start: float = word.get("start", bucket_start)
        w_end:   float = word.get("end",   w_start)

        if bucket:
            projected      = " ".join(bucket + [word_text])
            over_words     = len(bucket) >= max_words
            over_chars     = len(projected) > max_chars
            over_dur       = (w_end - bucket_start) > max_duration
            sent_break     = sentence_aware and bool(_SENTENCE_END.search(bucket[-1]))

            if over_words or over_chars or over_dur or sent_break:
                _flush()
                bucket_start = w_start

        if not bucket:
            bucket_start = w_start

        bucket.append(word_text)
        bucket_end = w_end

    _flush()
    return result


# ── Text-based formats ─────────────────────────────────────────────────────────

def export_txt(segments: list, include_timestamps: bool = False) -> str:
    lines = []
    for seg in segments:
        if include_timestamps:
            lines.append(f"{_readable_time(seg['start'])} {seg['text']}")
        else:
            lines.append(seg["text"])
    return "\n\n".join(lines)


def export_srt(segments: list) -> str:
    blocks = []
    for i, seg in enumerate(segments, start=1):
        block = (
            f"{i}\n"
            f"{_srt_time(seg['start'])} --> {_srt_time(seg['end'])}\n"
            f"{seg['text']}\n"
        )
        blocks.append(block)
    return "\n".join(blocks)


def export_vtt(segments: list) -> str:
    lines = ["WEBVTT", ""]
    for seg in segments:
        lines.append(f"{_vtt_time(seg['start'])} --> {_vtt_time(seg['end'])}")
        lines.append(seg["text"])
        lines.append("")
    return "\n".join(lines)


def export_csv(segments: list, include_end_time: bool = True) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    header = ["Start", "End", "Text"] if include_end_time else ["Start", "Text"]
    writer.writerow(header)
    for seg in segments:
        if include_end_time:
            writer.writerow([
                _readable_time(seg["start"]),
                _readable_time(seg["end"]),
                seg["text"],
            ])
        else:
            writer.writerow([_readable_time(seg["start"]), seg["text"]])
    return output.getvalue()


# ── Binary formats ─────────────────────────────────────────────────────────────

def export_docx(
    name: str,
    segments: list,
    include_timestamps: bool = True,
    save_path: str | None = None,
) -> bytes | None:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    title = doc.add_heading(name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Transcribed with LocalScribe · {len(segments)} segments")
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer

    for seg in segments:
        para = doc.add_paragraph()
        if include_timestamps:
            ts_run = para.add_run(_readable_time(seg["start"]) + " ")
            ts_run.font.color.rgb = RGBColor(130, 130, 200)
            ts_run.font.size = Pt(9)
        text_run = para.add_run(seg["text"])
        text_run.font.size = Pt(11)

    if save_path:
        doc.save(save_path)
        return None

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _TranscriptPDF(FPDF):
    def __init__(self, title: str, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.doc_title = title
        self.alias_nb_pages()

    def header(self):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(180, 160, 240)
        self.cell(0, 8, self.doc_title, align="L")
        self.set_text_color(150, 150, 150)
        self.cell(0, 8, "LocalScribe", align="R")
        self.ln(4)
        self.set_draw_color(60, 60, 90)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def export_pdf(
    name: str,
    segments: list,
    include_timestamps: bool = True,
    save_path: str | None = None,
) -> bytes | None:
    pdf = _TranscriptPDF(title=name, orientation="P", unit="mm", format="A4")
    pdf.set_margins(20, 25, 20)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(100, 100, 150)
    pdf.cell(0, 12, name, align="C")
    pdf.ln(14)

    for seg in segments:
        if include_timestamps:
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(130, 120, 180)
            pdf.cell(0, 5, _readable_time(seg["start"]))
            pdf.ln(5)

        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(50, 50, 50)
        # FPDF's built-in Helvetica only covers latin-1; replace unrepresentable
        # characters rather than crashing. Embedding a TTF would give full Unicode.
        text = seg["text"].encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 6, text)
        pdf.ln(3)

    if save_path:
        pdf.output(save_path)
        return None
    return bytes(pdf.output())


# ── Self-contained HTML share page ────────────────────────────────────────────


